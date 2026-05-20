from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "vystupy" / "sjednocene_docx"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(31, 87, 130)
DARK = RGBColor(29, 37, 45)
MUTED = RGBColor(90, 96, 102)
ACCENT = RGBColor(183, 76, 45)
FILL = "EEF3F7"
FILL2 = "F7F4EF"
BORDER = "CBD5DF"

TERMS = [
    ("Název klubu", "Rock'n'Wall Praha z.s."),
    ("Provozovatel stěny", "Boulder Area s.r.o. / HUDY Boulder Karlín"),
    ("Místo kroužků", "HUDY Boulder Karlín, Křižíkova 684/91a, Praha 8"),
    ("Hlavní evidence", "KIS"),
    ("Běžný provoz stěny", "HUDY / Memberzone / Member Pro"),
    ("Účetnictví", "LP soft"),
    ("Start změny", "1. 9. 2026, školní rok 2026/2027"),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = tbl_borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, headers, rows, widths=(1.7, 5.1)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    for index, header in enumerate(headers):
        cell_text(table.cell(0, index), header, True)
        shade(table.cell(0, index), FILL)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell_text(cells[index], str(value), index == 0 and len(headers) == 2)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
    doc.add_paragraph()


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    cell = table.cell(0, 0)
    shade(cell, FILL2)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def setup_doc(title, subtitle):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.4)
    sec.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    title_style = styles["Title"]
    title_style.font.name = "Arial"
    title_style.font.size = Pt(22)
    title_style.font.color.rgb = DARK
    title_style.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, ACCENT, 9, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = sec.header.paragraphs[0]
    header.text = "Rock'n'Wall Praha z.s."
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = sec.footer.paragraphs[0]
    footer.text = "Pracovní sjednocená verze | 20. 5. 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    doc.add_paragraph(title, style="Title")
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.italic = True
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(12)
    add_note(
        doc,
        "Stav dokumentu",
        "Pracovní text k dalším úpravám. Názvosloví je sjednocené s předchozími podklady, ale formulace lze dál zkracovat nebo upravovat podle finálního nastavení KIS, plateb a dokumentů.",
    )
    return doc


def add_terms(doc):
    doc.add_heading("Sjednocené názvosloví", level=1)
    add_table(doc, ["Položka", "Používaný tvar"], TERMS, (1.75, 5.05))


def save(doc, filename):
    path = OUT / filename
    doc.save(path)
    return path


def build_00():
    doc = setup_doc(
        "00 Zdroje a ověřené údaje",
        "Souhrn ověřených údajů pro web, dokumenty a komunikaci s HUDY Boulder Karlín.",
    )
    add_terms(doc)
    doc.add_heading("Ověřené údaje spolku", level=1)
    add_table(
        doc,
        ["Položka", "Údaj"],
        [
            ("Název", "Rock'n'Wall Praha z.s."),
            ("Právní forma", "spolek"),
            ("IČO", "29542707"),
            ("Spisová značka", "L 82082 vedená u Městského soudu v Praze"),
            ("Sídlo", "č.ev. 402, 267 11 Vráž"),
            ("Nejvyšší orgán", "členská schůze"),
            ("Statutární orgán", "předseda"),
            ("Způsob jednání", "Předseda jedná za spolek samostatně."),
            ("Datum založení podle stanov", "26. 4. 2026"),
            ("Usnesení o zápisu", "Praha, 11. 5. 2026"),
        ],
        (1.8, 5.0),
    )
    doc.add_heading("Účel spolku", level=1)
    doc.add_paragraph(
        "Podpora, rozvoj a organizace sportovní, tělovýchovné, vzdělávací a komunitní činnosti v oblasti lezení, boulderingu, horolezectví a souvisejících pohybových aktivit."
    )
    doc.add_heading("Předmět činnosti", level=1)
    doc.add_paragraph(
        "Organizace tréninků, kurzů, kroužků, soustředění, soutěží, výprav, přednášek a dalších sportovních, vzdělávacích a komunitních akcí."
    )
    doc.add_heading("Zdrojové dokumenty", level=1)
    bullets(
        doc,
        [
            r"C:\Dokumenty Dejvicka\Rock n Wall Praha\vystup_Rock_n_Wall_Praha_zalozeni_spolku\01_Stanovy_RocknWall_Praha_zs.docx",
            r"C:\Dokumenty Dejvicka\Rock n Wall Praha\vystup_Rock_n_Wall_Praha_zalozeni_spolku\02_Zapis_z_ustavujici_schuze_a_listina_pritomnych.docx",
            r"C:\Dokumenty Dejvicka\Rock n Wall Praha\vystup_Rock_n_Wall_Praha_zalozeni_spolku\priloha_1695741396_0_vyhoveni_usnesenim_Fj_182940_2026_MSPH.pdf",
        ],
    )
    doc.add_heading("Veřejné zdroje HUDY", level=1)
    bullets(
        doc,
        [
            "Dětské kroužky a kurzy: https://www.hudysteny.cz/boulderkarlin/krouzky-a-lekce/detske-krouzky/",
            "Provozní a bezpečnostní řád: https://www.hudysteny.cz/boulderkarlin/provozni-rad/",
            "Obchodní a platební podmínky: https://www.hudysteny.cz/boulderkarlin/obchodni-a-platebni-podminky/",
            "Souhlas ke kroužkům: https://www.hudysteny.cz/boulderkarlin/wp-content/uploads/2023/07/HUDY-Boulder-Karlin_Souhlas-krouzky.pdf",
        ],
    )
    add_note(
        doc,
        "Důležité rozlišení",
        "Boulder Area s.r.o. zůstává provozovatelem stěny a běžného provozu HUDY Boulder Karlín. Rock'n'Wall Praha z.s. má od 1. 9. 2026 organizovat pravidelné dětské kroužky a klubovou činnost. Registrace, evidence, členství, docházka a platby za klubové kroužky mají být vedené přes KIS a účet klubu.",
    )
    return save(doc, "00_zdroje_a_overene_udaje_RocknWall_Praha.docx")


def build_01():
    doc = setup_doc(
        "01 Klubový web - struktura a finální texty",
        "Texty pro jednoduchou statickou microsite rocknwallpraha.cz.",
    )
    add_terms(doc)
    doc.add_heading("Cíl webu", level=1)
    doc.add_paragraph(
        "Web Rock'n'Wall Praha z.s. má být jednoduchá, profesionální a snadno spravovatelná microsite. Nemá nahrazovat KIS ani vytvářet další databázi osobních údajů."
    )
    doc.add_paragraph(
        "Hlavní úkol webu je informovat rodiče, vysvětlit přihlášení, odkázat do KIS, zveřejnit dokumenty a oddělit klubové kroužky od běžného provozu HUDY / Memberzone."
    )
    doc.add_heading("Úvodní stránka", level=1)
    add_table(
        doc,
        ["Část", "Finální text"],
        [
            ("H1", "Klubové lezecké kroužky v HUDY Boulder Karlín"),
            (
                "Perex",
                "Pravidelné dětské kroužky a pololetní kurzy organizuje Rock'n'Wall Praha z.s. ve spolupráci s HUDY Boulder Karlín. Evidence, členství, platby a docházka budou vedené v klubovém systému KIS.",
            ),
            (
                "Důležité upozornění",
                "Běžné vstupy, kredit a veřejný provoz stěny zůstávají v systému HUDY / Memberzone. Klubové kroužky od 1. 9. 2026 budou řešené přes Rock'n'Wall Praha z.s. a KIS.",
            ),
            ("Tlačítka", "Jak se přihlásit | Přehled kroužků | Dokumenty pro rodiče"),
        ],
        (1.55, 5.25),
    )
    doc.add_heading("Struktura webu", level=1)
    bullets(doc, ["Úvod", "Kroužky", "Přihlášení", "Platby", "Členství", "Pro rodiče", "Dokumenty", "Novinky", "Kontakt"])
    doc.add_heading("Kroužky", level=1)
    doc.add_paragraph(
        "Aktuální skupiny, kapacity a přihlášky budou vedené v KIS. Na webu klubu bude uveden základní přehled typů kroužků, pravidel a odkaz na registraci."
    )
    add_table(
        doc,
        ["Typ skupiny", "Popis"],
        [
            ("Začátečníci", "Pro děti, které s lezením začínají nebo chtějí pravidelný bezpečný pohyb na stěně."),
            ("Pokročilí", "Pro děti s předchozí zkušeností, které chtějí zlepšovat techniku, sílu a samostatnost."),
            ("Tréninkové skupiny", "Pro děti se zájmem o systematičtější trénink, závody a dlouhodobý sportovní růst."),
        ],
        (1.7, 5.1),
    )
    doc.add_heading("Přihlášení", level=1)
    numbers(
        doc,
        [
            "Rodič si na webu vybere vhodný typ kroužku a přejde do registračního formuláře v KIS.",
            "V KIS vyplní údaje dítěte, zákonného zástupce a potřebné souhlasy.",
            "Klub zkontroluje přihlášku a zařadí dítě do skupiny nebo na čekací listinu.",
            "Po zařazení obdrží rodič platební údaje a variabilní symbol.",
            "Po úhradě je dítě vedené jako aktivní člen a účastník kurzu.",
            "Trenér následně vede docházku v KIS.",
        ],
    )
    doc.add_heading("Platby", level=1)
    doc.add_paragraph(
        "Platby za pravidelné klubové kroužky nebudou probíhat přes kredit v Memberzone. Po přihlášení a zařazení dítěte do skupiny obdrží rodič platební údaje z KIS. Platba se hradí přímo na účet Rock'n'Wall Praha z.s."
    )
    add_table(
        doc,
        ["Položka", "Text"],
        [
            ("Příjemce", "Rock'n'Wall Praha z.s."),
            ("Účet klubu", "Bude doplněn."),
            ("Variabilní symbol", "Podle pokynu z KIS."),
            ("Účetnictví", "LP soft."),
        ],
        (1.55, 5.25),
    )
    doc.add_heading("Členství", level=1)
    doc.add_paragraph(
        "Dítě přijaté do pravidelného klubového kroužku bude členem Rock'n'Wall Praha z.s. Členství umožňuje vést přehlednou klubovou evidenci, organizovat sportovní činnost a připravit klub na budoucí podporu, závody a dotace."
    )
    doc.add_heading("Pro rodiče", level=1)
    add_table(
        doc,
        ["Téma", "Text"],
        [
            ("Před lekcí", "Dítě přichází včas, ve sportovním oblečení a s pitím. Podrobnosti k vybavení a půjčování lezeček budou doplněny před spuštěním registrací."),
            ("Bezpečnost", "V prostoru HUDY Boulder Karlín platí provozní a bezpečnostní pravidla sportoviště. Klubové kroužky se zároveň řídí pravidly klubu."),
            ("Omluvy", "Omluvy a docházka budou řešené přes KIS nebo podle pokynů koordinátora kroužků."),
            ("Kontakt", "Pro dotazy ke kroužkům bude sloužit klubový kontakt. Pro běžný provoz stěny zůstává kontaktem HUDY Boulder Karlín."),
        ],
        (1.55, 5.25),
    )
    doc.add_heading("Dokumenty", level=1)
    bullets(
        doc,
        [
            "Stanovy spolku",
            "Informace o zpracování osobních údajů",
            "Pravidla klubových kroužků",
            "Platební a storno podmínky",
            "Souhlas zákonného zástupce",
            "Souhlas / nesouhlas s fotografiemi a videem",
        ],
    )
    doc.add_heading("Kontakt", level=1)
    add_table(
        doc,
        ["Položka", "Údaj"],
        [
            ("Název", "Rock'n'Wall Praha z.s."),
            ("IČO", "29542707"),
            ("Sídlo", "č.ev. 402, 267 11 Vráž"),
            ("Místo kroužků", "HUDY Boulder Karlín, Křižíkova 684/91a, Praha 8"),
            ("E-mail", "Bude doplněn."),
        ],
        (1.55, 5.25),
    )
    return save(doc, "01_klubovy_web_struktura_a_texty.docx")


def build_02():
    doc = setup_doc(
        "02 Texty pro HUDY web",
        "Návrh úprav stránky dětských kroužků a souvisejících podmínek HUDY Boulder Karlín.",
    )
    add_terms(doc)
    doc.add_heading("Cíl úprav", level=1)
    doc.add_paragraph(
        "Na webu HUDY je potřeba jasně vysvětlit novou spolupráci s Rock'n'Wall Praha z.s. a oddělit běžný provoz stěny od klubových kroužků, registrací a plateb přes KIS."
    )
    doc.add_heading("Krátký box na začátek stránky Dětské kroužky", level=1)
    doc.add_paragraph("Dětské kroužky od školního roku 2026/2027")
    doc.add_paragraph(
        "Od 1. 9. 2026 budou pravidelné dětské lezecké kroužky a pololetní kurzy v HUDY Boulder Karlín organizované ve spolupráci s klubem Rock'n'Wall Praha z.s."
    )
    doc.add_paragraph(
        "HUDY Boulder Karlín zůstává místem konání a provozovatelem sportoviště. Klub Rock'n'Wall Praha z.s. zajišťuje registrace dětí do kroužků, členskou evidenci, platby, docházku a komunikaci ke klubové činnosti."
    )
    doc.add_paragraph(
        "Přihlašování do pravidelných dětských kroužků bude probíhat přes klubový systém KIS. Platby za klubové kroužky budou hrazeny přímo na účet Rock'n'Wall Praha z.s.; nebudou probíhat přes kredit v systému Memberzone."
    )
    doc.add_heading("Delší vysvětlení spolupráce", level=1)
    doc.add_paragraph(
        "Od školního roku 2026/2027 přechází organizace pravidelných dětských kroužků a pololetních kurzů pod sportovní klub Rock'n'Wall Praha z.s."
    )
    doc.add_heading("Co se mění", level=2)
    bullets(
        doc,
        [
            "Přihlášení do pravidelných dětských kroužků bude probíhat přes klubový systém KIS.",
            "Dítě bude po přijetí vedeno jako člen klubu.",
            "Platba za kroužek bude probíhat přímo na účet klubu Rock'n'Wall Praha z.s.",
            "Docházka, evidence a komunikace ke kroužkům budou vedeny v klubovém systému.",
        ],
    )
    doc.add_heading("Co zůstává", level=2)
    bullets(
        doc,
        [
            "Kroužky budou probíhat v prostorech HUDY Boulder Karlín.",
            "Pro pohyb v areálu stále platí provozní a bezpečnostní řád HUDY Boulder Karlín.",
            "Veřejné vstupy, kredit a běžný provoz stěny zůstávají v systému HUDY / Memberzone.",
        ],
    )
    doc.add_heading("Platební upozornění", level=1)
    doc.add_paragraph(
        "Platby za pravidelné dětské kroužky organizované klubem Rock'n'Wall Praha z.s. od 1. 9. 2026 neprobíhají přes kredit v systému Memberzone."
    )
    doc.add_paragraph(
        "Po přihlášení a zařazení dítěte do kroužku obdrží rodič platební údaje z klubového systému KIS. Platba se hradí přímo na účet klubu Rock'n'Wall Praha z.s. podle uvedeného variabilního symbolu."
    )
    doc.add_paragraph(
        "Platby za běžné vstupy, veřejné lekce, kredit, půjčovnu a další služby HUDY Boulder Karlín zůstávají beze změny podle obchodních a platebních podmínek provozovatele Boulder Area s.r.o."
    )
    doc.add_heading("Doporučené doplnění provozních pravidel", level=1)
    add_note(
        doc,
        "Text",
        "Pravidelné dětské kroužky organizované klubem Rock'n'Wall Praha z.s. se vedle provozního a bezpečnostního řádu HUDY Boulder Karlín řídí také pravidly a dokumenty klubu.",
    )
    doc.add_heading("Doporučené doplnění obchodních a platebních podmínek", level=1)
    add_note(
        doc,
        "Text",
        "Tyto obchodní a platební podmínky se vztahují na služby provozovatele Boulder Area s.r.o. Platby za pravidelné dětské kroužky organizované klubem Rock'n'Wall Praha z.s. probíhají přímo mezi zákonným zástupcem účastníka a klubem podle podmínek Rock'n'Wall Praha z.s.",
    )
    doc.add_heading("Co odstranit nebo přepsat ze stávající stránky", level=1)
    bullets(
        doc,
        [
            "Návod na vytvoření registrace dítěti v Memberzone.",
            "Přihlášení přes záložku „Celé kroužky a kurzy“ v Memberzone.",
            "Platbu nahráním kreditu.",
            "Přihlašování jednotlivých lekcí přes Memberzone, pokud se netýká jiné veřejné služby HUDY.",
        ],
    )
    return save(doc, "02_texty_pro_HUDY_web.docx")


def build_03():
    doc = setup_doc(
        "03 Workflow, data a administrativa",
        "Návrh provozního toku rodič -> registrace -> členství -> platba -> evidence -> docházka.",
    )
    add_terms(doc)
    doc.add_heading("Hlavní princip", level=1)
    doc.add_paragraph(
        "Web pouze informuje a směruje rodiče do KIS. KIS je hlavní evidence. Účetnictví se vede v LP soft. Platby za kroužky jdou přímo na účet klubu."
    )
    doc.add_heading("Workflow registrace", level=1)
    numbers(
        doc,
        [
            "Rodič přijde na HUDY web nebo klubový web.",
            "HUDY web vysvětlí novou spolupráci a odkazuje na klub.",
            "Klubový web podrobně vysvětlí kroužky, členství, platby a dokumenty.",
            "Rodič klikne na registraci do KIS.",
            "Rodič vyplní údaje dítěte a zákonného zástupce.",
            "Rodič potvrdí souhlasy.",
            "Klub zkontroluje přihlášení.",
            "Klub zařadí dítě do skupiny nebo dá přihlášce stav čekatel.",
            "KIS připraví platební údaje.",
            "Rodič zaplatí na účet klubu.",
            "Platba se spáruje nebo ručně označí.",
            "Dítě je aktivní člen a účastník kurzu.",
            "Trenér vede docházku.",
        ],
    )
    doc.add_heading("Stavy přihlášky", level=1)
    bullets(doc, ["nová", "čeká na kontrolu", "čeká na doplnění", "přijato", "čekatel", "odmítnuto / není kapacita", "čeká na platbu", "aktivní", "zrušeno", "archivováno"])
    doc.add_heading("Minimální datový model v KIS", level=1)
    add_table(
        doc,
        ["Entita", "Doporučená pole"],
        [
            ("Dítě / člen", "jméno, příjmení, datum narození, adresa, zákonný zástupce, zdravotní omezení, členství, kurz, platba, docházka, poznámka"),
            ("Zákonný zástupce", "jméno, příjmení, e-mail, telefon, adresa, vztah k dítěti, fakturační údaje, souhlasy"),
            ("Kurz / skupina", "název, pololetí, den a čas, věková skupina, úroveň, trenér, kapacita, cena, místo, stav"),
            ("Platba", "dítě, kurz, částka, splatnost, variabilní symbol, stav, datum přijetí, vazba na LP soft"),
            ("Docházka", "datum lekce, kurz, dítě, přítomen / omluven / neomluven, trenér, poznámka"),
            ("Dokument / souhlas", "typ dokumentu, verze, datum souhlasu, zákonný zástupce, vazba na dítě, archivace"),
        ],
        (1.65, 5.15),
    )
    doc.add_heading("Vazba na LP soft", level=1)
    doc.add_paragraph(
        "LP soft by měl zůstat účetní evidencí. KIS eviduje předpis platby a stav dítěte. Banka přijme platbu, platba se spáruje nebo označí v KIS a souhrn plateb se pravidelně kontroluje proti LP soft."
    )
    doc.add_heading("Připravenost na dotace NSA", level=1)
    bullets(doc, ["počet aktivních dětí", "věkové složení", "docházka", "trenéři", "pravidelná sportovní činnost", "závody a akce", "členské příspěvky", "dokumenty o členské základně", "základní hospodaření"])
    return save(doc, "03_workflow_data_administrativa.docx")


def build_04():
    doc = setup_doc(
        "04 Checklist přípravy",
        "Pracovní kontrolní seznam pro spuštění klubových kroužků od 1. 9. 2026.",
    )
    add_terms(doc)
    sections = [
        ("Web", ["potvrdit doménu klubu", "založit GitHub repozitář", "zapnout GitHub Pages", "připravit dokumenty ke stažení", "doplnit odkaz na KIS", "doplnit kontakt a účet klubu"]),
        ("Texty pro HUDY web", ["krátký oznam o spolupráci", "přepis přihlášení z Memberzone na KIS", "přepis plateb z kreditu na účet klubu", "odkaz na klubový web", "doplnění informace do provozních pravidel", "doplnění informace do obchodních a platebních podmínek"]),
        ("KIS", ["nastavit klub / organizaci", "nastavit role", "nastavit členskou evidenci", "nastavit kurzy / skupiny", "nastavit přihlašovací formulář", "nastavit stavy přihlášek", "nastavit předpisy plateb", "nastavit docházku", "nastavit exporty"]),
        ("Platby", ["potvrdit číslo účtu klubu", "rozhodnout strukturu variabilních symbolů", "nastavit splatnosti", "nastavit kontrolu plateb", "nastavit proces upomínek", "nastavit kontrolu proti LP soft"]),
        ("Dokumenty", ["stanovy pro veřejné zveřejnění", "GDPR informace", "pravidla klubových kroužků", "platební a storno podmínky", "souhlas zákonného zástupce", "souhlas s fotografiemi a videem", "postup při úrazu"]),
        ("Před spuštěním", ["testovací přihláška v KIS", "test platebního předpisu", "test spárování platby", "test docházky trenéra", "kontrola textů na HUDY webu", "kontrola odkazů z HUDY na klub", "kontrola odkazů z klubu do KIS"]),
    ]
    for heading, items in sections:
        doc.add_heading(heading, level=1)
        bullets(doc, ["☐ " + item for item in items])
    doc.add_heading("Pracovní harmonogram", level=1)
    add_table(
        doc,
        ["Období", "Úkoly"],
        [
            ("Květen-červen 2026", "webová struktura, dokumenty, KIS, texty pro HUDY, bankovní účet, LP soft"),
            ("Červenec 2026", "test registrace, plateb, docházky, dokumentů a komunikace rodičům"),
            ("Srpen 2026", "spuštění informací a registrací, kontrola kapacit, platby, přístupy trenérů"),
            ("Od 1. září 2026", "ostrý provoz, docházka v KIS, platby na účet klubu, průběžná administrativa"),
        ],
        (1.8, 5.0),
    )
    return save(doc, "04_checklist_pripravy.docx")


def main():
    paths = [build_00(), build_01(), build_02(), build_03(), build_04()]
    for path in paths:
        with ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        if any(marker in xml for marker in ["D?tsk?", "P?ihl", "Krou?k", "?lenstv", "?koln"]):
            raise RuntimeError(f"Document appears to contain broken Czech encoding: {path}")
        if "ě" not in xml and "ř" not in xml:
            raise RuntimeError(f"Document does not appear to contain Czech diacritics: {path}")
        print(path)


if __name__ == "__main__":
    main()
